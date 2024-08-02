from flask import Flask, request, render_template
from docx import Document
import pdfkit
import git

app = Flask(__name__)

@app.route('/edit')
def edit_docx():
    doc = Document('./Resume Updated.docx')
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    text = '\n'.join(full_text)
    return render_template('editor.html', text=text)

@app.route('/save', methods=['POST'])
def save_docx():
    new_text = request.form['text']
    doc = Document()
    for line in new_text.split('\n'):
        doc.add_paragraph(line)
    doc.save('./Resume Updated.docx')

    # Convert DOCX to PDF
    pdfkit.from_file('./Resume Updated.docx', './resume.pdf')

    # Push the updated PDF to GitHub
    repo = git.Repo('/app')
    repo.git.add('resume.pdf')
    repo.git.commit('-m', 'Updated resume.pdf')
    origin = repo.remote(name='origin')
    origin.push()

    return "File updated and pushed to GitHub"

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5204)
